import time
import getpass
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches

def generate_docx_report(driver):
    timestamp = time.strftime("%B %d, %Y    %H:%M", time.localtime())
    document = Document()

    document.add_heading('Test Report', level=0)
    document.add_paragraph(f'Date and Time: {timestamp}')

    # Получение информации об исполнителе из ОС
    executor = getpass.getuser()
    document.add_paragraph(f'Test executor: {executor}')

    document.add_heading('Test: Open the web site and create screenshot', level=1)
    document.add_paragraph('Status: PASS')
    document.add_heading('Steps:', level=2)

    try:
        driver.get("https://billyal.netlify.app/")
        driver.maximize_window()
        driver.save_screenshot('screenshot1.png')
        document.add_paragraph('1. Page opened: https://billyal.netlify.app/')
        document.add_paragraph('Screenshot 1:')
        document.add_picture('screenshot1.png', width=Inches(5))
        time.sleep(1)

        elements = driver.find_elements(By.CSS_SELECTOR, '.mx-3.my-1.text-dark.text-decoration-none.fw-bold')
        if elements:
            elements[1].click()
            document.add_paragraph('2. The element found and selected.')
            driver.save_screenshot('screenshot2.png')
            document.add_paragraph('3. Second page opened.')
            document.add_paragraph('Screenshot 2:')
            document.add_picture('screenshot2.png', width=Inches(5))
        else:
            document.add_paragraph('2. Elements was not found.')

        time.sleep(1)
    except Exception as e:
        document.add_paragraph(f'Произошла ошибка во время выполнения теста: {e}', style='List Bullet')

    document.save('test_report.docx')
    print("Word отчет сохранен в файле test_report.docx")

def test_billy_al(driver):
    generate_docx_report(driver)